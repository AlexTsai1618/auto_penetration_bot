from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import numpy as np
from pandas.compat import StringIO

memfile = StringIO()
t = np.arange(0.0, 2.0, 0.01)
s = 1 + np.sin(2*np.pi*t)
plt.plot(t, s)
plt.savefig(memfile)

document = Document()
document.add_heading('Report',0)
document.add_picture(memfile, width=Inches(1.25))

document.save('report.docx')
memfile.close()